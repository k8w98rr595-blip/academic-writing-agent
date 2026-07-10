from __future__ import annotations

import io
import re
import zipfile
from pathlib import Path

from docx import Document as WordDocument
from fastapi import HTTPException, status

from .text import normalize_text


MAX_DOCX_BYTES = 5 * 1024 * 1024
MAX_ZIP_ENTRIES = 1000
MAX_UNCOMPRESSED_BYTES = 25 * 1024 * 1024
DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def validate_docx_upload(filename: str, content_type: str, payload: bytes) -> None:
    safe_name = Path(filename or "").name
    if not safe_name.lower().endswith(".docx") or safe_name != filename:
        raise HTTPException(status_code=status.HTTP_422_UNPROCESSABLE_CONTENT, detail="Only safe .docx filenames are accepted")
    if content_type and content_type not in {DOCX_MIME, "application/octet-stream"}:
        raise HTTPException(status_code=status.HTTP_422_UNPROCESSABLE_CONTENT, detail="Invalid Word MIME type")
    if len(payload) < 4 or len(payload) > MAX_DOCX_BYTES or payload[:2] != b"PK":
        raise HTTPException(status_code=status.HTTP_422_UNPROCESSABLE_CONTENT, detail="Invalid or oversized Word document")
    try:
        with zipfile.ZipFile(io.BytesIO(payload)) as archive:
            infos = archive.infolist()
            if len(infos) > MAX_ZIP_ENTRIES or sum(item.file_size for item in infos) > MAX_UNCOMPRESSED_BYTES:
                raise HTTPException(status_code=status.HTTP_422_UNPROCESSABLE_CONTENT, detail="Word document expands beyond safe limits")
            names = {item.filename for item in infos}
            if any(
                name.startswith(("/", "\\"))
                or "\\" in name
                or ".." in Path(name).parts
                for name in names
            ):
                raise HTTPException(status_code=status.HTTP_422_UNPROCESSABLE_CONTENT, detail="Unsafe Word archive path")
            if "word/document.xml" not in names or any(name.lower().endswith("vbaproject.bin") for name in names):
                raise HTTPException(status_code=status.HTTP_422_UNPROCESSABLE_CONTENT, detail="Unsupported Word document structure")
            for name in names:
                if name.endswith(".rels"):
                    relation_data = archive.read(name)
                    if re.search(br"TargetMode\s*=\s*[\"']External[\"']", relation_data, flags=re.IGNORECASE):
                        raise HTTPException(status_code=status.HTTP_422_UNPROCESSABLE_CONTENT, detail="External Word relationships are not accepted")
    except zipfile.BadZipFile as error:
        raise HTTPException(status_code=status.HTTP_422_UNPROCESSABLE_CONTENT, detail="Invalid Word archive") from error


def extract_docx_text(payload: bytes) -> str:
    try:
        document = WordDocument(io.BytesIO(payload))
    except Exception as error:
        raise HTTPException(status_code=status.HTTP_422_UNPROCESSABLE_CONTENT, detail="Unable to parse Word document") from error
    blocks = [normalize_text(paragraph.text) for paragraph in document.paragraphs if normalize_text(paragraph.text)]
    for table in document.tables:
        rows = []
        for row in table.rows:
            cells = [normalize_text(cell.text) for cell in row.cells]
            if any(cells):
                rows.append(" | ".join(cells))
        if rows:
            blocks.append("\n".join(rows))
    return "\n\n".join(blocks)


def build_docx(title: str, paragraphs: list[dict], version_number: int) -> bytes:
    document = WordDocument()
    document.core_properties.title = title
    document.core_properties.subject = f"Paperlight version {version_number}"
    for paragraph in paragraphs:
        document.add_paragraph(str(paragraph["text"]))
    stream = io.BytesIO()
    document.save(stream)
    return stream.getvalue()
