from __future__ import annotations

import io
import zipfile
from dataclasses import replace

from fastapi.testclient import TestClient
from docx import Document as WordDocument

from services.api.app.documents import DOCX_MIME, validate_docx_upload
from services.api.app.database import session_scope
from services.api.app.models import AnalysisRun
from services.api.app.service import new_id
from services.api.app.text import assert_protected_equal
from services.api.app import main as main_module


def create_document(client: TestClient, headers: dict[str, str], text: str) -> dict:
    response = client.post(
        "/api/v1/documents",
        headers=headers,
        data={"title": "Ethics of Data Reuse", "text": text},
    )
    assert response.status_code == 201, response.text
    return response.json()["document"]


def test_document_word_count_boundary(client: TestClient, headers: dict[str, str]):
    too_short = client.post(
        "/api/v1/documents",
        headers=headers,
        data={"title": "Too short", "text": "word " * 499},
    )
    assert too_short.status_code == 422
    assert too_short.json()["detail"] == "Paper must contain at least 500 words"

    minimum = client.post(
        "/api/v1/documents",
        headers=headers,
        data={"title": "Minimum length", "text": "word " * 500},
    )
    assert minimum.status_code == 201, minimum.text
    assert minimum.json()["document"]["currentVersion"]["wordCount"] == 500


def test_owner_only_mock_workflow(client: TestClient, headers: dict[str, str], coursework_text: str):
    assert client.get("/api/v1/documents").status_code == 401
    document = create_document(client, headers, coursework_text)

    analysis = client.post(f"/api/v1/documents/{document['id']}/analyses", headers=headers)
    assert analysis.status_code == 201, analysis.text
    result = analysis.json()["analysis"]["result"]
    assert result["isMock"] is True
    assert result["provider"] == "Mock Pangram"
    assert result["providerModel"] == "mock"
    assert result["providerModelVersion"] == "mock-pangram-4-shape-v1"
    assert result["disclaimer"].startswith("This is a probabilistic AI writing risk signal")
    assert round(result["aiGeneratedPercent"] + result["aiAssistedPercent"] + result["humanPercent"], 1) == 100.0
    assert all(row["classification"] in {"ai_generated", "ai_assisted"} for row in result["spans"])
    assert "qualityChecks" not in result
    job = client.get(f"/api/v1/jobs/{analysis.json()['jobId']}/events", headers=headers)
    assert job.status_code == 200
    assert '"status": "completed"' in job.text

    rewrite = client.post(
        f"/api/v1/documents/{document['id']}/rewrite-sessions",
        headers=headers,
        json={"version_id": document["currentVersion"]["id"]},
    )
    assert rewrite.status_code == 201
    paragraph = document["currentVersion"]["paragraphs"][1]
    patch_response = client.post(
        f"/api/v1/rewrite-sessions/{rewrite.json()['rewriteSession']['id']}/messages",
        headers=headers,
        json={"instruction": "Make this more direct", "paragraph_id": paragraph["id"], "selected_text": ""},
    )
    assert patch_response.status_code == 201, patch_response.text
    patch = patch_response.json()["patch"]
    assert patch["isMock"] is True
    assert patch["originalText"] != patch["revisedText"]
    assert patch["rewriteSessionId"] == rewrite.json()["rewriteSession"]["id"]
    assert patch["revisionNumber"] == 1
    assert patch["contextScope"] == "selection"

    refined_response = client.post(
        f"/api/v1/rewrite-sessions/{rewrite.json()['rewriteSession']['id']}/messages",
        headers=headers,
        json={
            "instruction": "Make the current suggestion more concise",
            "paragraph_id": paragraph["id"],
            "selected_text": "",
            "previous_patch_id": patch["id"],
            "context_scope": "section",
        },
    )
    assert refined_response.status_code == 201, refined_response.text
    refined = refined_response.json()["patch"]
    assert refined["revisionNumber"] == 2
    assert refined["supersedesPatchId"] == patch["id"]
    assert refined["originalText"] == patch["originalText"]
    assert refined["revisedText"] != patch["revisedText"]
    refreshed = client.get(f"/api/v1/documents/{document['id']}", headers=headers).json()["document"]
    patch_states = {item["id"]: item["status"] for item in refreshed["patches"]}
    assert patch_states[patch["id"]] == "superseded"
    assert patch_states[refined["id"]] == "pending"
    assert next(item for item in refreshed["patches"] if item["id"] == refined["id"])["revisionNumber"] == 2

    accepted = client.post(
        f"/api/v1/patches/{refined['id']}/accept",
        headers=headers,
        json={"expected_base_version_id": refined["baseVersionId"]},
    )
    assert accepted.status_code == 200, accepted.text
    assert accepted.json()["document"]["currentVersion"]["number"] == 2
    assert accepted.json()["document"]["analysis"]["isStale"] is True
    assert client.post(
        f"/api/v1/patches/{patch['id']}/accept",
        headers=headers,
        json={"expected_base_version_id": patch["baseVersionId"]},
    ).status_code == 409

    reanalysis = client.post(f"/api/v1/documents/{document['id']}/analyses", headers=headers)
    assert reanalysis.status_code == 201, reanalysis.text
    reanalysis_result = reanalysis.json()["analysis"]["result"]
    assert reanalysis_result["riskComparison"]["beforePercent"] == result["combinedRiskPercent"]
    assert reanalysis_result["riskComparison"]["afterPercent"] == reanalysis_result["combinedRiskPercent"]

    version_one = next(row for row in accepted.json()["document"]["versions"] if row["number"] == 1)
    restored = client.post(
        f"/api/v1/documents/{document['id']}/versions/{version_one['id']}/restore",
        headers=headers,
        json={"expected_current_version_id": accepted.json()["document"]["currentVersion"]["id"]},
    )
    assert restored.status_code == 200
    assert restored.json()["document"]["currentVersion"]["number"] == 3
    assert restored.json()["document"]["currentVersion"]["source"] == "restore"

    exported = client.post(f"/api/v1/documents/{document['id']}/exports", headers=headers)
    assert exported.status_code == 200
    assert exported.content[:2] == b"PK"
    assert exported.headers["content-type"].startswith(DOCX_MIME)
    exported_word = WordDocument(io.BytesIO(exported.content))
    assert exported_word.paragraphs[0].text == "Introduction"
    assert any("ethical data reuse" in paragraph.text for paragraph in exported_word.paragraphs)

    deleted = client.delete(f"/api/v1/documents/{document['id']}", headers=headers)
    assert deleted.status_code == 204
    assert client.get(f"/api/v1/documents/{document['id']}", headers=headers).status_code == 404


def test_first_pass_mock_returns_reviewable_preview_without_changing_version(
    client: TestClient, headers: dict[str, str], coursework_text: str
):
    document = create_document(client, headers, coursework_text)
    assert client.post(f"/api/v1/documents/{document['id']}/analyses", headers=headers).status_code == 201
    response = client.post(
        f"/api/v1/documents/{document['id']}/first-pass-rewrite",
        headers=headers,
        json={"version_id": document["currentVersion"]["id"]},
    )
    assert response.status_code == 201, response.text
    payload = response.json()
    assert payload["applied"] is False
    assert payload["patch"]["isMock"] is True
    assert payload["patch"]["status"] == "pending"
    assert payload["targetParagraphCount"] >= 1
    refreshed = client.get(f"/api/v1/documents/{document['id']}", headers=headers).json()["document"]
    assert refreshed["currentVersion"]["id"] == document["currentVersion"]["id"]
    assert refreshed["analysis"]["isStale"] is False


def test_first_pass_real_mode_applies_all_safe_revisions_as_one_version(
    client: TestClient, headers: dict[str, str], coursework_text: str, monkeypatch
):
    document = create_document(client, headers, coursework_text)
    assert client.post(f"/api/v1/documents/{document['id']}/analyses", headers=headers).status_code == 201

    async def fake_first_pass(passages, **_):
        return {
            "revisions": [
                {
                    **passage,
                    "revisedText": passage["originalText"].replace("It is important to note that", "Evidence indicates that", 1),
                    "reason": "Removed formulaic framing.",
                }
                for passage in passages
            ],
            "isMock": False,
            "provider": "DeepSeek",
            "modelVersion": "deepseek-v4-pro",
            "validatorModelVersion": "deepseek-v4-flash",
        }

    monkeypatch.setattr(main_module, "settings", replace(main_module.settings, rewrite_mode="deepseek"))
    monkeypatch.setattr(main_module, "reserve_provider_calls", lambda *_args, **_kwargs: {})
    monkeypatch.setattr(main_module, "propose_first_pass_rewrites", fake_first_pass)
    response = client.post(
        f"/api/v1/documents/{document['id']}/first-pass-rewrite",
        headers=headers,
        json={"version_id": document["currentVersion"]["id"]},
    )
    assert response.status_code == 201, response.text
    payload = response.json()
    assert payload["applied"] is True
    assert payload["document"]["currentVersion"]["number"] == 2
    assert payload["document"]["currentVersion"]["source"] == "agent-first-pass"
    assert payload["document"]["analysis"]["isStale"] is True
    assert payload["revisedParagraphCount"] == payload["targetParagraphCount"]
    accepted = [patch for patch in payload["document"]["patches"] if patch["status"] == "accepted"]
    assert len(accepted) == payload["revisedParagraphCount"]


def test_stale_manual_save_is_rejected(client: TestClient, headers: dict[str, str], coursework_text: str):
    document = create_document(client, headers, coursework_text)
    paragraphs = document["currentVersion"]["paragraphs"]
    payload = {"base_version_id": document["currentVersion"]["id"], "paragraphs": paragraphs}
    first = client.patch(f"/api/v1/documents/{document['id']}", headers=headers, json=payload)
    assert first.status_code == 200
    second = client.patch(f"/api/v1/documents/{document['id']}", headers=headers, json=payload)
    assert second.status_code == 409


def test_legacy_analysis_record_remains_read_only_and_unchanged(
    client: TestClient, headers: dict[str, str], coursework_text: str
):
    document = create_document(client, headers, coursework_text)
    legacy_result = {
        "estimate": 42.0,
        "isMock": True,
        "fusionStatus": "provider-agreement",
        "providers": [{"name": "Historical Provider"}],
        "spans": [{"paragraphId": document["currentVersion"]["paragraphs"][0]["id"], "start": 0, "end": 5}],
        "disclaimer": "Historical probabilistic result.",
    }
    with session_scope() as db:
        db.add(
            AnalysisRun(
                id=new_id("analysis"),
                document_id=document["id"],
                version_id=document["currentVersion"]["id"],
                status="completed",
                provider_mode="legacy-dual",
                result=legacy_result,
            )
        )
    response = client.get(f"/api/v1/documents/{document['id']}", headers=headers)
    assert response.status_code == 200
    assert response.json()["document"]["analysis"]["result"] == legacy_result


def test_protected_tokens_cannot_change():
    try:
        assert_protected_equal("The result was 42% [3].", "The result was 41% [3].")
    except Exception as error:
        assert getattr(error, "status_code", None) == 409
    else:
        raise AssertionError("Protected-token change was accepted")


def test_docx_external_relationship_is_rejected():
    stream = io.BytesIO()
    with zipfile.ZipFile(stream, "w") as archive:
        archive.writestr("word/document.xml", "<document />")
        archive.writestr("word/_rels/document.xml.rels", '<Relationship TargetMode="External" Target="https://example.com"/>')
    try:
        validate_docx_upload("paper.docx", DOCX_MIME, stream.getvalue())
    except Exception as error:
        assert getattr(error, "status_code", None) == 422
    else:
        raise AssertionError("External Word relationship was accepted")


def test_safe_docx_import_is_editable_and_deleted(
    client: TestClient, headers: dict[str, str], coursework_text: str
):
    word = WordDocument()
    word.add_paragraph(coursework_text)
    stream = io.BytesIO()
    word.save(stream)
    response = client.post(
        "/api/v1/documents",
        headers=headers,
        data={"title": "Imported paper"},
        files={"file": ("paper.docx", stream.getvalue(), DOCX_MIME)},
    )
    assert response.status_code == 201, response.text
    document = response.json()["document"]
    assert document["currentVersion"]["source"] == "docx"
    assert document["currentVersion"]["paragraphs"]
    assert client.delete(f"/api/v1/documents/{document['id']}", headers=headers).status_code == 204
